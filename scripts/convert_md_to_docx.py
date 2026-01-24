#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
MarkdownファイルをWord文書（.docx）に変換するスクリプト
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

def convert_markdown_to_docx(md_file_path, output_path):
    """
    MarkdownファイルをWord文書に変換
    """
    # Markdownファイルを読み込み
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Word文書を作成
    doc = Document()

    # スタイル設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10.5)

    # 行を分割
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 空行
        if not line:
            i += 1
            continue

        # H1 (# )
        if line.startswith('# '):
            text = line[2:].strip()
            paragraph = doc.add_heading(text, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # H2 (## )
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)

        # H3 (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)

        # H4 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)

        # 箇条書き (- )
        elif line.startswith('- '):
            text = line[2:].strip()
            # **太字**を処理
            text = process_inline_formatting(text)
            paragraph = doc.add_paragraph(style='List Bullet')
            add_formatted_text(paragraph, text)

        # 区切り線 (---)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)

        # テーブル (| で始まる行)
        elif line.startswith('|'):
            # テーブルを処理
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # 1つ戻す

            if len(table_lines) > 2:  # ヘッダー + 区切り + データ
                add_table_to_doc(doc, table_lines)

        # 【】で囲まれた見出し
        elif line.startswith('【') and '】' in line:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(11)

        # 通常の段落
        else:
            text = process_inline_formatting(line)
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, text)

        i += 1

    # 保存
    doc.save(output_path)
    print(f'変換完了: {output_path}')

def process_inline_formatting(text):
    """
    インライン形式（**太字**、リンクなど）を処理
    """
    return text

def add_formatted_text(paragraph, text):
    """
    太字やリンクを含むテキストをパラグラフに追加
    """
    # **太字**を処理
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 太字
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif '[' in part and '](' in part and ')' in part:
            # リンク処理（簡易版）
            link_match = re.search(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                before = part[:link_match.start()]
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                after = part[link_match.end():]

                if before:
                    paragraph.add_run(before)

                # リンクテキスト + URLを表示
                run = paragraph.add_run(f'{link_text} ({link_url})')
                run.font.color.rgb = RGBColor(0, 0, 255)

                if after:
                    paragraph.add_run(after)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def add_table_to_doc(doc, table_lines):
    """
    MarkdownテーブルをWord文書に追加
    """
    # ヘッダーとデータを分離
    header_line = table_lines[0]
    # 区切り線をスキップ
    data_lines = table_lines[2:]

    # セルを分割
    header_cells = [cell.strip() for cell in header_line.split('|')[1:-1]]

    # テーブル作成
    table = doc.add_table(rows=1 + len(data_lines), cols=len(header_cells))
    table.style = 'Light Grid Accent 1'

    # ヘッダー行
    header_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = header_row.cells[i]
        # **を削除
        cell_text = cell_text.replace('**', '')
        cell.text = cell_text
        # 太字にする
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # データ行
    for row_idx, data_line in enumerate(data_lines):
        data_cells = [cell.strip() for cell in data_line.split('|')[1:-1]]
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(data_cells):
            # **を削除
            cell_text = cell_text.replace('**', '')
            row.cells[col_idx].text = cell_text

if __name__ == '__main__':
    # 職務経歴書のパス
    md_file = r'C:\Users\tetsu\dev\energy-env\learning_memos\職務経歴書.md'
    output_file = r'C:\Users\tetsu\dev\energy-env\learning_memos\職務経歴書.docx'

    convert_markdown_to_docx(md_file, output_file)
